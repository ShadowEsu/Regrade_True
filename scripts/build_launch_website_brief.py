from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "deliverables" / "Regrade_Launch_Website_Brief.docx"

NAVY = "0C2144"
BLUE = "1769E0"
PALE_BLUE = "EEF5FF"
INK = "182230"
MUTED = "5D6B7E"
LINE = "DDE5F0"
GREEN = "18794E"
AMBER = "9A6700"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_cell_border(cell, color=LINE, val='single', sz='6'):
    tcPr = cell._tc.get_or_add_tcPr(); borders = tcPr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ['top','left','bottom','right']:
        tag = qn(f'w:{edge}'); e = borders.find(tag)
        if e is None:
            e = OxmlElement(f'w:{edge}'); borders.append(e)
        e.set(qn('w:val'), val); e.set(qn('w:sz'), sz); e.set(qn('w:color'), color)

def set_table_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for i, width in enumerate(widths):
        if i < len(grid.gridCol_lst): grid.gridCol_lst[i].w = Twips(width)
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells): row.cells[i].width = Twips(width)

def add_run(p, text, size=11, color=INK, bold=False, italic=False):
    r = p.add_run(text); r.font.name = 'Aptos'; r._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos'); r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
    r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color); r.bold = bold; r.italic = italic
    return r

def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    for margin in ['top_margin','right_margin','bottom_margin','left_margin']:
        setattr(sec, margin, Inches(0.82))
    sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)
    styles = doc.styles
    normal = styles['Normal']; normal.font.name='Aptos'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.16
    for name, size, color, before, after in [('Title',30,NAVY,0,6),('Heading 1',18,NAVY,18,7),('Heading 2',13,BLUE,12,5),('Heading 3',11,NAVY,8,3)]:
        s=styles[name]; s.font.name='Aptos Display' if name!='Heading 3' else 'Aptos'; s._element.rPr.rFonts.set(qn('w:ascii'),s.font.name); s._element.rPr.rFonts.set(qn('w:hAnsi'),s.font.name); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    if 'Kicker' not in styles:
        s=styles.add_style('Kicker', WD_STYLE_TYPE.PARAGRAPH); s.font.name='Aptos'; s._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos'); s.font.size=Pt(9); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(BLUE); s.paragraph_format.space_after=Pt(6)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; add_run(footer, 'RE GRADE  •  LAUNCH WEBSITE BRIEF', 8, MUTED, True)

def p(doc, text='', style=None, align=None, after=None):
    para=doc.add_paragraph(style=style)
    if align is not None: para.alignment=align
    if text: add_run(para, text)
    if after is not None: para.paragraph_format.space_after=Pt(after)
    return para

def bullet(doc, text):
    para=doc.add_paragraph(style='List Bullet'); para.paragraph_format.space_after=Pt(3); add_run(para, text, 10.2)
    return para

def callout(doc, label, text, fill=PALE_BLUE, color=BLUE):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; set_table_widths(t,[9360])
    c=t.cell(0,0); set_cell_shading(c,fill); set_cell_border(c, fill, 'nil', '0'); set_cell_margins(c,160,200,160,200)
    q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(2); add_run(q,label.upper(),8.5,color,True)
    q=c.add_paragraph(); q.paragraph_format.space_after=Pt(0); add_run(q,text,10.4,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def section_title(doc, kicker, title, intro=None):
    p(doc,kicker.upper(),'Kicker'); p(doc,title,'Heading 1')
    if intro:
        x=p(doc); add_run(x,intro,11,MUTED); x.paragraph_format.space_after=Pt(8)

def add_feature_card(doc, title, headline, body, bullets, website, status):
    t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT; set_table_widths(t,[1900,7460])
    for c in t.rows[0].cells: set_cell_margins(c,150,170,150,170); set_cell_border(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(t.cell(0,0),PALE_BLUE)
    a=t.cell(0,0).paragraphs[0]; add_run(a,title.upper(),8.5,BLUE,True); a.add_run('\n'); add_run(a,status,9,NAVY,True)
    c=t.cell(0,1); x=c.paragraphs[0]; add_run(x,headline,12,NAVY,True)
    x=c.add_paragraph(); add_run(x,body,10.1,INK)
    x=c.add_paragraph(); add_run(x,'Website idea: ',9.3,BLUE,True); add_run(x,website,9.3,MUTED)
    for item in bullets:
        x=c.add_paragraph(style='List Bullet'); x.paragraph_format.space_after=Pt(1); add_run(x,item,9.4,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def add_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; set_table_widths(t,widths)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell_shading(c,'E8EEF5'); set_cell_border(c); set_cell_margins(c); q=c.paragraphs[0]; add_run(q,h,9,NAVY,True)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; set_cell_border(c); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP; q=c.paragraphs[0]; add_run(q,val,9.1,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def page_break(doc): doc.add_page_break()

def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc=Document(); style_doc(doc)
    # Cover
    p(doc,'RE GRADE','Kicker',WD_ALIGN_PARAGRAPH.CENTER)
    p(doc,'Launch Website\nMessaging & Design Brief','Title',WD_ALIGN_PARAGRAPH.CENTER)
    sub=p(doc,align=WD_ALIGN_PARAGRAPH.CENTER); add_run(sub,'A professional, evidence-first student support platform for understanding grades, learning from marked exams, and making stronger appeals.',14,MUTED)
    sub.paragraph_format.space_after=Pt(20)
    callout(doc,'Core positioning','Understand your grade. Learn from it. Make the strongest appeal.',PALE_BLUE,BLUE)
    p(doc,'Prepared for Regrade launch website and product demo',align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
    p(doc,'July 2026  •  Working product brief',align=WD_ALIGN_PARAGRAPH.CENTER,after=0)
    page_break(doc)

    section_title(doc,'01  •  Product story','Regrade is more than an appeal tool—without losing the appeal.',"The launch site should make one thing clear immediately: Regrade helps students inspect graded work, understand what happened, prepare a respectful evidence-led appeal when appropriate, and turn marked exams into better preparation for the next one.")
    p(doc,'Hero copy','Heading 2')
    callout(doc,'Hero headline','Know what happened on your grade. Know what to do next.',PALE_BLUE,BLUE)
    p(doc,'Supporting copy: Upload a marked exam, PDF, or screenshot. Regrade organizes the visible score, rubric, and teacher feedback; Mr Whale helps you understand the result, practise the weak spots, and prepare a clear appeal when the evidence supports one.')
    p(doc,'Primary CTA: Review a marked exam     Secondary CTA: See how it works',after=12)
    p(doc,'Brand voice','Heading 2')
    bullet(doc,'Calm, precise, and credible—not combative, childish, or overly legal.')
    bullet(doc,'Student-first, but never promises a grade change or assumes instructor error.')
    bullet(doc,'Professional enough for colleges and school partnerships; warm enough for stressed students.')
    bullet(doc,'Mr Whale is the guide, not the entire product. The interface stays clean and academic.')
    p(doc,'The four-part value proposition','Heading 2')
    add_table(doc,['Pillar','What students get','Website framing'],[
        ['Review','A clear read of visible marks, rubric rows, and feedback.','“See the grade, not just the number.”'],
        ['Appeal','A factual, respectful draft when a concrete discrepancy exists.','“Ask the right question with the right evidence.”'],
        ['Study','Exam-only patterns and check-off preparation before finals.','“Turn past marks into your next study plan.”'],
        ['Coach','Mr Whale explains a mark in plain language and stays in context.','“A calm guide beside the work.”'],
    ],[1500,3700,4160])

    section_title(doc,'02  •  Website structure','A persuasive page that stays easy to scan.',"Use a restrained editorial layout: white or pale-gray canvas, navy ink, one confident Regrade blue, generous spacing, and very limited glass or gradients. The whale can animate subtly inside product moments, not across every section.")
    add_table(doc,['Section','Purpose','Design idea'],[
        ['1. Hero','State the promise and give an immediate product proof point.','Split layout: concise copy left; Review Studio interface preview right.'],
        ['2. How it works','Reduce anxiety around the first upload.','Three numbered steps with real-looking document/rubric/coach miniatures.'],
        ['3. Product proof','Show Review, Appeal, Study, and Coach in context.','Horizontal story with one active product screen at a time.'],
        ['4. Trust','Make privacy, evidence, and student control concrete.','Quiet policy-style cards: “You control uploads”, “No invented findings”, “Review before sending.”'],
        ['5. Educators & families','Show future partnership path without overpromising.','Consent-first supervisor card and institution-ready analytics roadmap.'],
        ['6. Final CTA','Return to the student outcome.','“Bring one marked exam. Leave with a clearer next step.”'],
    ],[1300,3100,4960])

    page_break(doc)
    section_title(doc,'03  •  Feature messaging','What the launch site should show—and what it should say.')
    add_feature_card(doc,'Appeal','Evidence before escalation.','Regrade reviews visible grading evidence and helps students request clarification or prepare an appeal that is specific, respectful, and grounded in the assignment.',['Reads visible scores, deductions, rubric items, and teacher comments.','Flags possible score/rubric mismatches as questions to review—not accusations.','Creates an editable professor-safe draft; students remain in control of sending.'],'Show a before/after: a confusing marked PDF becomes three traceable questions and a clean draft.','Available now')
    add_feature_card(doc,'Study','Finals Prep from actual marked exams.','Study starts empty. Once a student analyzes a marked exam, Regrade groups recurring deductions, rubric misses, and teacher feedback into a personal review plan.',['Uses exams only; excludes homework and unsupported assumptions.','Prioritizes recurring patterns across marked work.','Saves check-offs to the student profile while retaining evidence links.'],'Use a calm checklist with “2 exams” evidence badges and the original question feedback underneath.','Available now')
    add_feature_card(doc,'Review Studio','A document-led learning workspace.','The new Study Review Studio connects question-level evidence to practical next steps and places Mr Whale beside the mark, not in a disconnected chatbot.',['Green: visible strength. Blue: practise this skill. Amber: clarify the mark.','Every annotation is tied to extracted score, rubric, or teacher feedback.','Mr Whale can answer questions in the context of the selected mark.'],'Desktop product shot: evidence cards left, compact Mr Whale conversation right, no noisy dashboard chrome.','Available now — evidence map')
    add_feature_card(doc,'Mr Whale','The academic guide, not a generic bot.','Mr Whale explains a marked question, helps students revise, and helps shape a thoughtful appeal. The response appears as a short, readable speech bubble with a restrained typing animation.',['Plain-language explanations of marks and rubrics.','One focused follow-up at a time.','Never fabricates policy, teacher intent, or a guaranteed result.'],'Use the pixel whale as a small transparent mascot beside the response. Keep the page professional; avoid large cartoon panels.','Available now')
    add_feature_card(doc,'Connections','Bring in the work students already use.','Students can search for a learning platform, connect it when supported, or upload a marked file directly. Manual upload is always available.',['Featured first: Google Classroom, Canvas, Moodle.','Search, not a long dropdown.','Connection status is honest; unavailable sync is never presented as live.'],'A single search field with three familiar initial results and a quiet “Connect later” option.','Partially available')

    page_break(doc)
    section_title(doc,'04  •  Trust and quality','The parts of the story that make Regrade partnership-ready.')
    p(doc,'Evidence-first AI','Heading 2')
    bullet(doc,'Treats files, rubrics, and student notes as evidence—not instructions.')
    bullet(doc,'Separates visible marks, literal rubric support, and uncertainty.')
    bullet(doc,'Reads handwriting conservatively; unclear marks are flagged, never guessed.')
    bullet(doc,'Does not profile a professor from a single exam or claim bias/bad faith.')
    bullet(doc,'Defaults to clarification when evidence does not support a concrete correction.')
    p(doc,'Recommended launch-site trust copy','Heading 2')
    callout(doc,'Trust block','Regrade does not decide whether a grade is “fair.” It helps you read the marked work, identify what is visible, and ask a clearer question. You review every draft before it goes anywhere.', 'F6F8FB', NAVY)
    p(doc,'Privacy and account control','Heading 2')
    bullet(doc,'Students own their account, marked-work analysis, Study checklist, and connected platforms.')
    bullet(doc,'AI settings and analysis alerts are controlled in profile.')
    bullet(doc,'Account deletion is visible in Settings & Account.')
    bullet(doc,'Future supervisor access is consent-first: the learner accepts, chooses shared data, and can revoke access.')
    p(doc,'A professional launch site must avoid','Heading 2')
    bullet(doc,'“Win back your points” guarantees, aggressive claims, or “teachers make mistakes” as the core message.')
    bullet(doc,'Pretending that every LMS automatically imports grades or that a parent can access a child’s account without consent.')
    bullet(doc,'Claims that AI handwriting, grading, or instructor behavior is infallible.')

    section_title(doc,'05  •  Educators, parents, and institutional partnerships','A second audience with stricter boundaries.',"The supervisor workspace is designed for a parent or educator helping a learner—not for bypassing the learner’s account. It should be presented as a partnership path, not a feature that is already fully live.")
    p(doc,'Parent / supervisor','Heading 2')
    p(doc,'Support a learner while they remain in control. The future workspace should allow an invitation, learner acceptance, selected sharing scopes, approved-evidence review, and suggested drafts for learner approval.')
    p(doc,'Educator','Heading 2')
    p(doc,'Find learning gaps without replacing professional judgment. With course authorization, educators can view question difficulty, common rubric gaps, and approved cohort summaries.')
    p(doc,'Institution','Heading 2')
    p(doc,'Offer a trustworthy evidence and student-support layer with consent controls, retention policy, approved integrations, and a clear audit trail.')
    callout(doc,'Class report requirement','Class mean, median, distribution, pass rate, and question difficulty must come from an institution-authorized gradebook connection and teacher/course permission. Regrade must not infer class statistics from a student upload.', 'FFF7E6', AMBER)

    section_title(doc,'06  •  Product roadmap for the website','Be transparent about what is live, next, and institution-dependent.')
    add_table(doc,['Capability','Launch status','What is required next'],[
        ['Appeal analysis, drafting, Coach','Live product direction','Deploy production AI API, test with de-identified marked-work evaluation set.'],
        ['Study checklist and Review Studio evidence map','Live product direction','Continue improving question-level extraction and student corrections.'],
        ['True page overlays on original PDF','Next build','Opt-in page retention, private storage, page-coordinate annotations, manual correction tools, retention/deletion policy.'],
        ['Automatic LMS import / Auto Mode','Institution-dependent','OAuth/Developer Keys, server sync jobs, encrypted token storage, user approval before any draft.'],
        ['Cohort report','Institution-dependent','Teacher role, roster/course authorization, official gradebook import, privacy review, aggregation rules.'],
        ['Supervisor learner links','Next build','One-time invitations, learner acceptance, selected sharing scopes, revocation, audit log, email delivery.'],
        ['Pets','Later / optional','Opt-in companion only; never competes with the assessment interface.'],
    ],[2350,1750,5460])

    section_title(doc,'07  •  Demo story','A five-minute launch demo that proves the promise.')
    for i, step in enumerate([
        ('Start with one marked exam','A student uploads a marked PDF, image, or screenshot with a visible score, rubric, or teacher feedback.'),
        ('Show the evidence read','Regrade identifies question-level marks, comments, and uncertainty. Make clear: unclear handwriting is flagged, not guessed.'),
        ('Open Study Review Studio','Select an exam from the Study plan. Show the colored evidence map and Mr Whale beside the selected question.'),
        ('Ask the learning question','Ask: “What should I practise before the final?” Mr Whale gives a focused explanation tied to the mark.'),
        ('Show the appeal boundary','Open a potential issue and show the clarification/appeal draft. The student edits it; Regrade never sends it automatically.'),
        ('Close with the system','Study has captured the recurring pattern, while Appeal remains instantly available whenever evidence supports it.'),
    ], 1):
        x=doc.add_paragraph(); x.paragraph_format.space_before=Pt(5); x.paragraph_format.space_after=Pt(4); add_run(x,f'{i}. ',11,BLUE,True); add_run(x,step[0]+': ',11,NAVY,True); add_run(x,step[1],10.5,INK)

    section_title(doc,'08  •  Final website copy bank','Short lines ready to test on the launch site.')
    for label, value in [
        ('Hero','Understand your grade. Learn from it. Make the strongest appeal.'),
        ('Upload CTA','Bring one marked exam.'),
        ('Study CTA','Turn your marked exams into a better finals plan.'),
        ('Appeal CTA','Ask the right question with the right evidence.'),
        ('Coach CTA','Ask Mr Whale about the mark.'),
        ('Trust line','Clear evidence. Honest uncertainty. You stay in control.'),
        ('Supervisor line','Support a learner without taking over their account.'),
    ]:
        x=doc.add_paragraph(); x.paragraph_format.space_after=Pt(4); add_run(x,label.upper()+"  ",8.5,BLUE,True); add_run(x,value,11,NAVY,True)
    p(doc,'Recommended footer note', 'Heading 2')
    p(doc,'Regrade provides educational support, not legal or institutional decision-making. Students should review original marked work and follow their school’s policy before submitting an appeal.')
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__': main()
